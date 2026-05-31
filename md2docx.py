#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Конвертация Markdown → DOCX с аккуратной вёрсткой (GFM-таблицы, отступы, шрифты).
Требуется pandoc в PATH. Постобработка: python-docx.

Запуск:
  python md2docx.py "C:\\path\\to\\file.md"
  python md2docx.py file.md -o "C:\\out\\custom.docx"
"""

from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor


def _find_pandoc() -> str | None:
    exe = shutil.which("pandoc")
    if exe:
        return exe
    candidates = [
        Path(r"C:\Program Files\Pandoc\pandoc.exe"),
        Path(r"C:\Program Files (x86)\Pandoc\pandoc.exe"),
        Path.home() / r"AppData\Local\Pandoc\pandoc.exe",
    ]
    for p in candidates:
        if p.is_file():
            return str(p)
    return None


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for child in list(tr_pr):
        if child.tag == qn("w:tblHeader"):
            tr_pr.remove(child)
    tr_pr.append(OxmlElement("w:tblHeader"))


def _set_table_borders(table, color: str = "BFBFBF", sz_eighth_pt: str = "8") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz_eighth_pt)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def _set_cell_margins(cell, top=80, left=120, bottom=80, right=120) -> None:
    """Margins in twips (1/20 pt)."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def beautify_docx(docx_path: Path) -> None:
    doc = Document(str(docx_path))

    body_style = doc.styles["Normal"]
    body_font = body_style.font
    body_font.name = "Calibri"
    body_font.size = Pt(11)
    body_font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    pf = body_style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    heading_sizes = {"Heading 1": Pt(20), "Heading 2": Pt(15), "Heading 3": Pt(13), "Heading 4": Pt(12)}
    for name, pt in heading_sizes.items():
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = "Calibri Light"
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        st.font.size = pt
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True

    for p in doc.paragraphs:
        name = p.style.name if p.style else ""
        if name in ("Source Code", "Verbatim"):
            for r in p.runs:
                r.font.name = "Consolas"
                r.font.size = Pt(9)
            pf = p.paragraph_format
            pf.left_indent = Cm(0.6)
            pf.right_indent = Cm(0.6)
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            shd = parse_xml(r'<w:shd {ns} w:fill="F3F3F3" w:val="clear"/>'.format(ns=nsdecls("w")))
            p_pr = p._p.get_or_add_pPr()
            p_pr.append(shd)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            table.style = "Table Grid"
        except (KeyError, ValueError):
            pass
        _set_table_borders(table)
        if table.rows:
            _mark_header_row(table.rows[0])
        for r_idx, row in enumerate(table.rows):
            is_header = r_idx == 0
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                _set_cell_margins(cell)
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    for r in p.runs:
                        r.font.name = "Calibri"
                        r.font.size = Pt(10)
                        if is_header:
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
                if is_header:
                    _set_cell_shading(cell, "D9E2F3")

    for p in doc.paragraphs:
        name = p.style.name if p.style else ""
        if name in ("Compact", "List Paragraph"):
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.15

    for p in doc.paragraphs:
        if p.style and p.style.name == "Title":
            p.paragraph_format.space_after = Pt(12)
            for r in p.runs:
                r.font.name = "Calibri Light"
                r.font.size = Pt(26)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.save(str(docx_path))


def run_pandoc(pandoc: str, md_path: Path, docx_path: Path) -> None:
    cmd = [
        pandoc,
        str(md_path),
        "-f",
        "gfm+smart+pipe_tables+hard_line_breaks",
        "-t",
        "docx",
        "-o",
        str(docx_path),
        "--standalone",
        "--toc",
        "--toc-depth=3",
        "--highlight-style=tango",
        "--metadata",
        f"title={md_path.stem}",
    ]
    subprocess.run(cmd, check=True, capture_output=True, text=True, encoding="utf-8", errors="replace")


def convert(md_path: Path, out_path: Path | None) -> Path:
    md_path = md_path.expanduser().resolve()
    if not md_path.is_file():
        raise FileNotFoundError(md_path)

    pandoc = _find_pandoc()
    if not pandoc:
        raise RuntimeError(
            "Не найден pandoc. Установите Pandoc и добавьте в PATH "
            "(https://pandoc.org/installing.html)."
        )

    if out_path is None:
        out_path = md_path.with_suffix(".docx")
    else:
        out_path = out_path.expanduser().resolve()
        if out_path.suffix.lower() != ".docx":
            out_path = out_path.with_suffix(".docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td) / "tmp.docx"
        run_pandoc(pandoc, md_path, tmp)
        shutil.copyfile(tmp, out_path)

    beautify_docx(out_path)
    return out_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Markdown → DOCX (pandoc + оформление)")
    parser.add_argument("markdown", type=Path, help="Путь к .md")
    parser.add_argument("-o", "--output", type=Path, default=None, help="Путь к .docx (по умолчанию рядом с .md)")
    args = parser.parse_args()

    try:
        out = convert(args.markdown, args.output)
    except subprocess.CalledProcessError as e:
        err = (e.stderr or e.stdout or "").strip()
        print("pandoc завершился с ошибкой:", e.returncode, file=sys.stderr)
        if err:
            print(err, file=sys.stderr)
        return 1
    except Exception as e:
        print(str(e), file=sys.stderr)
        return 1

    print(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
